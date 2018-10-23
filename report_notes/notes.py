import docx
from openpyxl import load_workbook
import re
from docx.shared import Pt

# 获取文档对象
file_path_notes = r"D:\data\附注\输出正文1018.docx"
file_notes = docx.Document("输出正文.docx")

# 第二次找到中间表
# file_path_mid = r"D:\data\test\科目余额表报表\1\输出中间表1.xlsx"
file_path_mid = r"D:\work\2中间表.xlsx"
wb_mid = load_workbook(filename=file_path_mid, data_only=True)
sheets_mid = wb_mid.sheetnames
sheet_first_mid = sheets_mid[0]  # 中间表1
ws_mid = wb_mid[sheet_first_mid]  # 中间表工作区1

# 中间附注表二
sheet_second_mid = sheets_mid[1]
ws_mid_second = wb_mid[sheet_second_mid]
# print("B806:", "%.2f" % ws_mid_second["B806"].value)

# 中间附注表三
sheet_three_mid = sheets_mid[2]
ws_mid_three = wb_mid[sheet_three_mid]

# ######构建一个需要替换的字典#######
num = 2
replace_dict = {}
for i in ws_mid_three["A"][1:]:
    key = str(i.value)
    replace_dict[key] = str(ws_mid_three["B"+str(num)].value)
    num += 1
# print("replace_dict:",replace_dict)
# ##########################附注替换开始############################
# 替换段落中的文字
for para in file_notes.paragraphs:
    for k in replace_dict:
        if k in para.text:
            para.text = para.text.replace(k, replace_dict[k])
# num_par=0
# for para in file_notes.paragraphs:
#
#     print(num_par, para.text)
#     print("***************")
#     num_par += 1
# company_name = str(ws_mid_three["B2"].value)
# audit_year = str(ws_mid_three["B3"].value)
par6 = file_notes.paragraphs[6].text
par8 = file_notes.paragraphs[8].text
par57 = file_notes.paragraphs[57].text
file_notes.paragraphs[6].text = ""
file_notes.paragraphs[8].text = ""
file_notes.paragraphs[57].text = ""

run6 = file_notes.paragraphs[6].add_run(par6)
run6.font.name = '宋体'
run6.font.size = Pt(16)
run6.font.bold = True

run8 = file_notes.paragraphs[8].add_run(par8)
run8.font.name = '宋体'
run8.font.size = Pt(16)
run8.font.bold = True

run57 = file_notes.paragraphs[57].add_run(par6)
run57.font.name = '黑体'
run57.font.size = Pt(16)
run57.font.bold = True

# 替换表格中的文字
# print(111,file_notes.tables[1].cell(5,2).text)
zzs = ws_mid_three["B12"].value
zzs = '%.f%%' % (zzs * 100)
sds = ws_mid_three["B13"].value
sds = '%.f%%' % (sds * 100)
# file_notes.tables[1].cell(1,2).text = zzs
# file_notes.tables[1].cell(5,2).text = sds
run1 = file_notes.tables[1].cell(1, 2).paragraphs[0].add_run(zzs)
run1.font.name = 'Arial Narrow'
run1.font.size = Pt(9)
run2 = file_notes.tables[1].cell(5, 2).paragraphs[0].add_run(sds)
run2.font.name = 'Arial Narrow'
run2.font.size = Pt(9)

# ######################开始往附注中填数###########################


def standard_num(x):
    """
    :param x:需要标准化的数字
    :return: 标准化后的数字
    """
    if x == "None":
        return "0.00"
    elif x == "0" or x == "0.00":
        return "0.00"
    else:
        if '.' in x:
            if x[-3] == '.':
                num = len(x) % 3
                if num == 0:
                    n = x[0:-3]
                elif num == 1:
                    n = "00" + x[0:-3]
                elif num == 2:
                    n = "0" + x[0:-3]
                lis = re.findall(r'.{3}', n)
                c = ','.join(lis)

                c = c + x[-3:]
            else:
                x = x + "0"
                num = len(x) % 3
                if num == 0:
                    n = x[0:-3]
                elif num == 1:
                    n = "00" + x[0:-3]
                elif num == 2:
                    n = "0" + x[0:-3]
                lis = re.findall(r'.{3}', n)
                c = ','.join(lis)
                c = c + x[-3:]

        else:
            num = len(x) % 3
            if num == 0:
                n = x
            elif num == 1:
                n = "00" + x
            elif num == 2:
                n = "0" + x

            lis = re.findall(r'.{3}', n)
            c = ','.join(lis)
            c = c + ".00"
        d = c.lstrip("0")
        if d.startswith("-") and d[1] == ",":
            d = d.replace(",", "", 1)
        return d


def fill_first_table_two(table_index, start_line, first_column_num,
                         first_column_index, second_column_num, second_column_index,
                         list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid[list_cell[i - start_line]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)
    for v in range(start_line, second_column_num):
        run = file_notes.tables[table_index].cell(v, second_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid[list_cell[first_column_num - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_first_table_two2(table_index, start_line, first_column_num,
                          first_column_index, second_column_num, second_column_index,
                          list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号的集合
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            str(ws_mid[list_cell[i - start_line]].value))
        run.font.name = '宋体'
        run.font.size = Pt(9)
    for v in range(start_line, second_column_num):
        run = file_notes.tables[table_index].cell(v, second_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid[list_cell[first_column_num - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_first_table_two_per(table_index, start_line, first_column_num,
                             first_column_index, second_column_num, second_column_index,
                             list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            dec_per(ws_mid[list_cell[i - start_line]].value))
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(9)

    for v in range(start_line, second_column_num):
        run = file_notes.tables[table_index].cell(v, second_column_index).paragraphs[0].add_run(
            dec_per(ws_mid[list_cell[first_column_num - start_line + (v - start_line)]].value))
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(9)


def fill_second_table_two(table_index, start_line, first_column_num,
                          first_column_index, second_column_num, second_column_index,
                          list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[i - start_line]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)
    for v in range(start_line, second_column_num):
        run = file_notes.tables[table_index].cell(v, second_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[first_column_num - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_second_table_two2(table_index, start_line, first_column_num,
                           first_column_index, second_column_num, second_column_index,
                           list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            str(ws_mid_second[list_cell[i - start_line]].value))
        run.font.name = '宋体'
        run.font.size = Pt(9)
    for v in range(start_line, second_column_num):
        run = file_notes.tables[table_index].cell(v, second_column_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[first_column_num - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_second_table_one(table_index, start_line, first_column_num,
                          first_column_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格的列索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            str(ws_mid_second[list_cell[i - start_line]].value))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_second_table_one_per(table_index, start_line, first_column_num,
                              first_column_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格的列索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            dec_per(ws_mid_second[list_cell[i - start_line]].value))
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(9)


def fill_first_table_one(table_index, start_line, first_column_num,
                         first_column_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格的列索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        run = file_notes.tables[table_index].cell(i, first_column_index).paragraphs[0].add_run(
            standard_num(str(ws_mid[list_cell[i - start_line]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_second_table_three(table_index, start_line, first_column, first_index,
                            second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        run = file_notes.tables[table_index].cell(i, first_index).paragraphs[0].add_run(
            str(ws_mid_second[list_cell[i - start_line]].value))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for v in range(start_line, second_column):
        run = file_notes.tables[table_index].cell(v, second_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[first_column - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for n in range(start_line, three_column):
        run = file_notes.tables[table_index].cell(n, three_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[
                        list_cell[first_column - start_line + second_column - start_line + (n - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_second_table_three2(table_index, start_line, first_column, first_index,
                             second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        run = file_notes.tables[table_index].cell(i, first_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[i - start_line]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for v in range(start_line, second_column):
        run = file_notes.tables[table_index].cell(v, second_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[list_cell[first_column - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for n in range(start_line, three_column):
        run = file_notes.tables[table_index].cell(n, three_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid_second[
                        list_cell[first_column - start_line + second_column - start_line + (n - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


def fill_first_table_three(table_index, start_line, first_column, first_index,
                           second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        run = file_notes.tables[table_index].cell(i, first_index).paragraphs[0].add_run(
            str(ws_mid[list_cell[i - start_line]].value))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for v in range(start_line, second_column):
        run = file_notes.tables[table_index].cell(v, second_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid[list_cell[first_column - start_line + (v - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)

    for n in range(start_line, three_column):

        run = file_notes.tables[table_index].cell(n, three_index).paragraphs[0].add_run(
            standard_num(
                str(ws_mid[list_cell[first_column - start_line + second_column - start_line + (n - start_line)]].value)))
        run.font.name = '宋体'
        run.font.size = Pt(9)


# 把小数变成百分数
def dec_per(rate):
    if rate != None:
        b = '%.2f%%' % (rate * 100)
    else:
        b = "0.00%"
    return b


# 处理特殊的字体格式的修改

def modify_font(table_index, y, x, object_cell):
    """
    :param table_index: 表格的索引
    :param x: 表中中单元格的横坐标
    :param y: 表中中单元格的纵坐标
    :param object_cell: 需要的修改的单元格
    :return:
    """
    run = file_notes.tables[table_index].cell(y, x).paragraphs[0].add_run(object_cell)
    run.font.name = '宋体'
    run.font.size = Pt(9)



###
# 一.货币资金
# file_notes.tables[2].style.font.name = u'宋体'
fill_first_table_two(2, 1, 5, 1, 5, 2, ["C8", "C9", "C10", "C11", "C14", "C15", "C16", "C17"])

# 二.以公允价值计量且其变动计入当期损益的金融资产
fill_first_table_two(3, 1, 5, 1, 5, 2, ["C23", "C24", "C25", "C28", "C34", "C35", "C36", "C39"])

modify_font(3, 6, 1, standard_num(str(ws_mid["C32"].value)))
modify_font(3, 6, 2, standard_num(str(ws_mid["C43"].value)))



# 三.应收票据及应收账款
fill_first_table_two(4, 1, 4, 1, 4, 2, ["C51", "C78", "AA78", "C55", "C106", "AA106"])

# 3.1应收票据
fill_first_table_two(5, 1, 4, 1, 4, 2, ["C49", "C50", "C51", "C53", "C54", "C55"])

# 3.2应收账款
# 取数部分
fill_first_table_two(6, 2, 7, 1, 7, 3, ["AA80", "AA81", "AA82", "AA83", "AA86",
                                        "AA108", "AA109", "AA110", "AA111", "AA114"])

# 取比例部分
fill_first_table_two_per(6, 2, 6, 2, 6, 4, ["AB80", "AB81", "AB82", "AB83",
                                            "AB108", "AB109", "AB110", "AB111"])

# # 3.3 期末余额前五名
# 取数部分
fill_second_table_two2(7, 1, 6, 0, 7, 1, ["A60", "A61", "A62", "A63", "A64",
                                          "B60", "B61", "B62", "B63", "B64", "B65"])
# 取比例部分
fill_second_table_one_per(7, 1, 7, 2, ["C60", "C61", "C62", "C63", "C64", "C65"])


# 四.预付款项
# 4.1预付款项
# 取数部分
fill_first_table_two(8, 2, 7, 1, 7, 3, ["C126", "C127", "C128", "C129", "C130",
                                        "C132", "C133", "C134", "C135", "C136"])

# 取比例部分
fill_first_table_two_per(8, 2, 6, 2, 6, 4, ["AB126", "AB127", "AB128", "AB129",
                                            "AB132", "AB133", "AB134", "AB135"])

# 4.3 期末余额前五大明细
# 取数部分
fill_second_table_two2(9, 1, 6, 0, 7, 1, ["A88", "A89", "A90", "A91", "A92",
                                          "B88", "B89", "B90", "B91", "B92", "B93"])

# 取比例部分
fill_second_table_one_per(9, 1, 7, 2, ["C88", "C89", "C90", "C91", "C92", "C93"])


# 五.其他应收款
# 5.1

modify_font(10, 1, 1, standard_num(str(ws_mid["C147"].value)))
modify_font(10, 2, 1, standard_num(str(ws_mid_second["B112"].value)))
modify_font(10, 3, 1, standard_num(str(ws_mid["AA179"].value)))
modify_font(10, 4, 1, standard_num(str(ws_mid["AA187"].value)))
modify_font(10, 1, 2, standard_num(str(ws_mid["C154"].value)))
modify_font(10, 2, 2, standard_num(str(ws_mid_second["C112"].value)))
modify_font(10, 3, 2, standard_num(str(ws_mid["AA207"].value)))
modify_font(10, 4, 2, standard_num(str(ws_mid["AA215"].value)))

# 5.2应收利息
fill_first_table_three(11, 1, 6, 0, 7, 1, 7, 2, ["A142", "A143", "A144", "A145", "A146",
                                                 "C142", "C143", "C144", "C145", "C146", "C147",
                                                 "C149", "C150", "C151", "C152", "C153", "C154"])

# # 5.3 应收股利
fill_second_table_three(12, 1, 5, 0, 6, 1, 6, 2, ["A108", "A109", "A110", "A111",
                                                  "B108", "B109", "B110", "B111", "B112",
                                                  "C108", "C109", "C110", "C111", "C112"])

# 5.4 其他应收款
# 5.4.1其他应收款(期末数)
# 取数部分
fill_first_table_two(13, 2, 7, 1, 7, 3, ["AA173", "AA174", "AA175", "AA176", "AA179",
                                         "AA201", "AA202", "AA203", "AA204", "AA207"])

# 取比例部分
fill_first_table_two_per(13, 2, 6, 2, 6, 4, ["AB173", "AB174", "AB175", "AB176",
                                             "AB201", "AB202", "AB203", "AB204"])

# # 5.5 其他应收款期末余额前五大
# 取数部分
fill_second_table_two2(14, 1, 6, 0, 7, 1, ["A173", "A174", "A175", "A176", "A177",
                                           "L173", "L174", "L175", "L176", "L177", "L178"])

# 取比例部分
fill_second_table_one_per(14, 1, 7, 2, ["E173", "E174", "E175", "E176", "E177", "E178"])

# 六.存货
fill_first_table_three(15, 1, 8, 0, 9, 1, 9, 2, ["A248", "A249", "A250", "A251", "A252", "A253", "A254",
                                                 "C248", "C249", "C250", "C251", "C252", "C253", "C254", "C255",
                                                 "C276", "C277", "C278", "C279", "C280", "C281", "C282", "C283"])

# 七.持有待售资产
# 拆分成3列和两列
fill_second_table_three(16, 1, 6, 0, 7, 1, 7, 2, ["A203", "A204", "A205", "A206", "A207",
                                                  "B203", "B204", "B205", "B206", "B207", "B208",
                                                  "C203", "C204", "C205", "C206", "C207", "C208"])

fill_second_table_two(16, 1, 7, 3, 6, 4, ["D203", "D204", "D205", "D206", "D207", "D208",
                                          "E203", "E204", "E205", "E206", "E207"])

# 八.一年内到期的非流动资产
fill_second_table_three(17, 1, 9, 0, 10, 1, 10, 2, ["A216", "A217", "A218", "A219", "A220", "A221", "A222", "A223",
                                                    "B216", "B217", "B218", "B219", "B220", "B221", "B222", "B223",
                                                    "B224",
                                                    "C216", "C217", "C218", "C219", "C220", "C221", "C222", "C223",
                                                    "C224"])

# 九.其他流动资产
fill_second_table_three(18, 1, 9, 0, 10, 1, 10, 2, ["A228", "A229", "A230", "A231", "A232", "A233", "A234", "A235",
                                                    "B228", "B229", "B230", "B231", "B232", "B233", "B234", "B235",
                                                    "B236",
                                                    "C228", "C229", "C230", "C231", "C232", "C233", "C234", "C235",
                                                    "C236"])

# 十.可供出售金融资产
modify_font(19, 1, 1, standard_num(str(ws_mid["C365"].value)))
modify_font(19, 4, 1, standard_num(str(ws_mid["C364"].value)))
modify_font(19, 8, 1, standard_num(str(ws_mid["C366"].value)))
modify_font(19, 1, 2, standard_num(str(ws_mid["C384"].value)))
modify_font(19, 4, 2, standard_num(str(ws_mid["C383"].value)))
modify_font(19, 8, 2, standard_num(str(ws_mid["C385"].value)))

# 十一.持有至到期投资
fill_second_table_three(20, 1, 9, 0, 10, 1, 10, 2, ["A275", "A229", "A230", "A231", "A232", "A233", "A234", "A235",
                                                    "B228", "B229", "B230", "B231", "B232", "B233", "B234", "B235",
                                                    "B236",
                                                    "C228", "C229", "C230", "C231", "C232", "C233", "C234", "C235",
                                                    "C236"])
# 十二.长期应收款
fill_first_table_two(21, 1, 9, 1, 9, 2, ["C414", "C415", "C416", "C417", "C418", "C419", "C420", "C421",
                                         "C451", "C452", "C453", "C454", "C455", "C456", "C457", "C458"])


modify_font(21, 5, 0, standard_num(str(ws_mid["C418"].value)))
modify_font(21, 6, 0, standard_num(str(ws_mid["C419"].value)))
modify_font(21, 7, 0, standard_num(str(ws_mid["C420"].value)))

# 十三.长期股权投资
# 期末数和期初数
fill_second_table_two(22, 1, 5, 1, 5, 4, ["B321", "B309", "B315", "B322", "C418",
                                          "K321", "K309", "K315", "K322", "K418"])

# 增加数和减少数
fill_second_table_two(22, 1, 5, 2, 5, 3, ["L321", "L309", "L315", "L322", "L418",
                                          "M321", "M309", "M315", "M322", "M418"])

# 十四.投资性房地产
# 拆分成2列和2列
fill_first_table_two(23, 1, 16, 1, 13, 2, ["C471", "C468", "C469", "C512", "C509", "C510", "AA512", "AA509",
                                           "AA510", "C543", "C540", "C541", "C571", "AA540", "AA541",
                                           "C476", "C473", "C474", "C517", "C514", "C515", "AA517", "AA514",
                                           "AA515", "C548", "C545", "C546"])

fill_first_table_two(23, 1, 13, 3, 16, 4, ["C496", "C493", "C494", "C527", "C524", "C525", "AA527", "AA524",
                                           "AA525", "C558", "C555", "C556",
                                           "C506", "C503", "C504", "C537", "C534", "C535", "AA537", "AA534",
                                           "AA535", "C568", "C565", "C566", "C570", "AA565", "AA566"])

# 十五.固定资产

# 拆分成2列和2列
fill_first_table_two(24, 1, 14, 1, 13, 2, ["C624", "C618", "C619", "C620", "C621", "C622", "C689", "C683", "C684",
                                           "C685", "C686", "C687", "C781",
                                           "C632", "C626", "C627", "C628", "C629", "C630", "C697", "C691", "C692",
                                           "C693", "C694", "C695"])

fill_first_table_two(24, 1, 13, 3, 14, 4, ["C664", "C658", "C659", "C660", "C661", "C662", "C713", "C707",
                                           "C708", "C709", "C710", "C711",
                                           "C680", "C674", "C675", "C676", "C677", "C678", "C729", "C723", "C724",
                                           "C725", "C726", "C727", "C780"])

# 十六.在建工程
fill_second_table_three(25, 2, 7, 0, 8, 1, 8, 2, ["A328", "A329", "A330", "A331", "A332",
                                                  "B328", "B329", "B330", "B331", "B332", "B333",
                                                  "C328", "C329", "C330", "C331", "C332", "C333"])

fill_second_table_two(25, 2, 8, 3, 8, 4, ["D328", "D329", "D330", "D331", "D332", "D333",
                                          "E328", "E329", "E330", "E331", "E332", "E333"])

fill_second_table_two(25, 2, 8, 5, 8, 6, ["F328", "F329", "F330", "F331", "F332", "F333",
                                          "G328", "G329", "G330", "G331", "G332", "G333"])

# 十七.固定资产清理
# 无

# 十八.无形资产
# 拆分成2列和两列
fill_first_table_two(27, 1, 6, 1, 6, 2, ["C885", "C883", "C880", "C881", "C882",
                                         "C892", "C890", "C887", "C888", "C889"])
fill_first_table_two(27, 1, 6, 3, 6, 4, ["C920", "C918", "C915", "C916", "C917",
                                         "C934", "C932", "C929", "C930", "C931"])
fill_first_table_two(27, 9, 14, 1, 14, 2, ["C942", "C940", "C937", "C938", "C939",
                                           "C949", "C947", "C944", "C945", "C946"])
fill_first_table_two(27, 9, 14, 3, 14, 4, ["C963", "C961", "C958", "C959", "C960",
                                           "C977", "C975", "C972", "C973", "C974"])

modify_font(27, 17, 1, standard_num(str(ws_mid["C1023"].value)))
modify_font(27, 17, 4, standard_num(str(ws_mid["C1022"].value)))

# 十九.长期待摊费用
fill_second_table_three(28, 1, 9, 0, 10, 1, 10, 2, ["C493", "C494", "C495", "C496", "C497", "C498", "C499", "C450",
                                                    "C493", "C494", "C495", "C496", "C497", "C498", "C499", "C500",
                                                    "C501",
                                                    "D493", "D494", "D495", "D496", "D497", "D498", "D499", "D500",
                                                    "D501"])
fill_second_table_three2(28, 1, 10, 3, 10, 4, 10, 5, ["E493", "E494", "E495", "E496", "E497", "E498", "E499", "E450",
                                                      "E501",
                                                      "F493", "F494", "F495", "F496", "F497", "F498", "F499", "F500",
                                                      "F501",
                                                      "G493", "G494", "G495", "G496", "G497", "G498", "G499", "G500",
                                                      "G501"])
# 二十.递延所得税资产、递延所得税负债
fill_second_table_one(29, 3, 9, 0, ["A508", "A509", "A510", "A511", "A512", "A513"])
fill_second_table_one(29, 11, 16, 0, ["A516", "A517", "A518", "A519", "A520"])

fill_second_table_two(29, 2, 17, 1, 17, 2, ["B507", "B508", "B509", "B510", "B511", "B512", "B513", "B514", "B515",
                                            "B516", "B517", "B518", "B519", "B520", "B521",
                                            "C507", "C508", "C509", "C510", "C511", "C512", "C513", "C514", "C515",
                                            "C516", "C517", "C518", "C519", "C520", "C521"
                                            ])

fill_second_table_two(29, 2, 17, 3, 17, 4, ["D507", "D508", "D509", "D510", "D511", "D512", "D513", "D514", "D515",
                                            "D516", "D517", "D518", "D519", "D520", "D521",
                                            "E507", "E508", "E509", "E510", "E511", "E512", "E513", "E514", "E515",
                                            "E516", "E517", "E518", "E519", "E520", "E521"
                                            ])

# 二十一.其他非流动资产
fill_second_table_three(30, 1, 7, 0, 8, 1, 8, 2, ["A546", "A547", "A548", "A549", "A550", "A551",
                                                  "B546", "B547", "B548", "B549", "B550", "B551", "B552",
                                                  "C546", "C547", "C548", "C549", "C550", "C551", "C552"])

# 二十二.短期借款
fill_first_table_two(31, 1, 6, 1, 6, 2, ["C1037", "C1038", "C1039", "C1040", "C1041",
                                         "C1043", "C1044", "C1045", "C1046", "C1047"])

# 二十三.应付票据及应付账款
fill_first_table_two(32, 1, 4, 1, 4, 2, ["C1072", "C1087", "AA1087", "C1076", "C1095", "AA1095"])

# 1.应付票据
fill_first_table_two(33, 1, 4, 1, 4, 2, ["C1070", "C1071", "C1072",
                                         "C1074", "C1075", "C1076"])

# 2.应付账款
# 取数部分
fill_first_table_two(34, 2, 7, 1, 7, 3, ["C1081", "C1082", "C1083", "C1084", "C1087",
                                         "C1089", "C1090", "C1091", "C1092", "C1095"])

# 取比例部分
fill_first_table_two_per(34, 2, 6, 2, 6, 4, ["AB1081", "AB1082", "AB1083", "AB1084",
                                             "AB1089", "AB1090", "AB1091", "AB1092"])

# 二十四.预收款项
# 取数部分
fill_first_table_two(35, 2, 7, 1, 7, 3, ["C1100", "C1101", "C1102", "C1103", "C1106",
                                         "C1108", "C1109", "C1110", "C1111", "C1114"])

# 取比例部分
fill_first_table_two_per(35, 2, 6, 2, 6, 4, ["AB1100", "AB1101", "AB1102", "AB1103",
                                             "AB1108", "AB1109", "AB1110", "AB1111"])

# 二十五.应付职工薪酬
fill_first_table_two(36, 1, 7, 1, 7, 2, ["C1160", "C1161", "C1162", "C1168", "C1169", "C1173",
                                         "C1175", "C1176", "C1177", "C1183", "C1184", "C1188"])

fill_first_table_two(36, 1, 7, 3, 7, 4, ["C1190", "C1191", "C1192", "C1198", "C1199", "C1203",
                                         "C1205", "C1206", "C1207", "C1213", "C1214", "C1218"])

# 二十六.应交税费
fill_first_table_two(37, 1, 13, 1, 13, 2, ["C1255", "C1256", "C1257", "C1258", "C1259", "C1260",
                                           "C1261", "C1262", "C1263", "C1264", "C1265", "C1266",
                                           "C1268", "C1269", "C1270", "C1271", "C1272", "C1273",
                                           "C1274", "C1275", "C1276", "C1277", "C1278", "C1279"])
# 二十七.其他应付款
fill_first_table_two(38, 1, 3, 1, 3, 2, ["C1289", "C1305", "C1297", "C1311"])
# file_notes.tables[38].cell(3, 1).text = standard_num(str(ws_mid_second["B608"].value))
# file_notes.tables[38].cell(3, 2).text = standard_num(str(ws_mid_second["C608"].value))
# file_notes.tables[38].cell(4, 1).text = standard_num(str(ws_mid["AA1305"].value))
# file_notes.tables[38].cell(4, 2).text = standard_num(str(ws_mid["AA1311"].value))

modify_font(38, 3, 1, standard_num(str(ws_mid_second["B608"].value)))
modify_font(38, 3, 2, standard_num(str(ws_mid_second["C608"].value)))
modify_font(38, 4, 1, standard_num(str(ws_mid["AA1305"].value)))
modify_font(38, 4, 2, standard_num(str(ws_mid["AA1311"].value)))

# 1.应付利息
fill_first_table_three(39, 1, 4, 0, 5, 1, 5, 2, ["A1283", "A1284", "A1285",
                                                 "C1283", "C1284", "C1285", "C1289",
                                                 "C1291", "C1292", "C1293", "C1297"])

# 2.应付股利
fill_first_table_three(40, 1, 3, 0, 4, 1, 4, 2, ["A1301", "A1302",
                                                 "C1301", "C1302", "C1305",
                                                 "C1307", "C1308", "C1311"])

# 3.其他应付款
# 3.1其他应付款分类
fill_second_table_three(41, 1, 8, 0, 9, 1, 9, 2, ["A601", "A602", "A603", "A604", "A605", "A606", "A607",
                                                  "B601", "B602", "B603", "B604", "B605", "B606", "B607", "B608",
                                                  "C601", "C602", "C603", "C604", "C605", "C606", "C607", "C608"])

# 3.2 账龄超过1年的大额其他应付款情况的说明
fill_first_table_three(42, 1, 6, 0, 7, 1, 6, 2, ["A611", "A612", "A613", "A614", "A615",
                                                 "B611", "B612", "B613", "B614", "B615", "B616",
                                                 "C611", "C612", "C613", "C614", "C615"])

# 二十八.一年内到期的非流动负债
fill_first_table_two(43, 1, 4, 1, 4, 2, ["C1319", "C1320", "C1321",
                                         "C1327", "C1328", "C1329"])

# file_notes.tables[43].cell(5, 1).text = standard_num(str(ws_mid_second["C1325"].value))
# file_notes.tables[43].cell(5, 2).text = standard_num(str(ws_mid_second["C1333"].value))
modify_font(43, 5, 1, standard_num(str(ws_mid_second["C1325"].value)))
modify_font(43, 5, 2, standard_num(str(ws_mid_second["C1333"].value)))


# 二十九.其他流动负债
fill_second_table_three(44, 1, 6, 0, 7, 1, 7, 2, ["A622", "A623", "A624", "A625", "A626",
                                                  "B622", "B623", "B624", "B625", "B626", "B627",
                                                  "C622", "C623", "C624", "C625", "C626", "C627"])

# 三十.长期借款
fill_first_table_two(45, 1, 6, 1, 6, 2, ["C1339", "C1340", "C1341", "C1342", "C1343",
                                         "C1351", "C1352", "C1353", "C1354", "C1355"])

# 三十一.长期应付款
fill_second_table_three(46, 1, 6, 0, 7, 1, 7, 2, ["A666", "A667", "A668", "A669", "A670",
                                                  "B666", "B667", "B668", "B669", "B670", "B671",
                                                  "C666", "C667", "C668", "C669", "C670", "C671"])

# 三十二.专项应付款
fill_second_table_three(47, 1, 8, 0, 9, 1, 9, 2, ["A675", "A676", "A677", "A678", "A679", "A680", "A681",
                                                  "B675", "B676", "B677", "B678", "B679", "B680", "B681", "B682",
                                                  "C675", "C676", "C677", "C678", "C679", "C680", "C681", "C682"])

fill_second_table_three2(47, 1, 9, 3, 9, 4, 8, 5, ["D675", "D676", "D677", "D678", "D679", "D680", "D681", "D682",
                                                   "E675", "E676", "E677", "E678", "E679", "E680", "E681", "E682",
                                                   "F675", "F676", "F677", "F678", "F679", "F680", "F681"])

# 三十三.预计负债
fill_first_table_two(48, 1, 8, 1, 8, 2, ["C1448", "C1449", "C1450", "C1451", "C1452", "C1453", "C1454",
                                         "C1456", "C1457", "C1458", "C1459", "C1460", "C1461", "C1462"])
fill_first_table_two(48, 1, 8, 3, 8, 4, ["C1464", "C1465", "C1466", "C1467", "C1468", "C1469", "C1470",
                                         "C1472", "C1473", "C1474", "C1475", "C1476", "C1477", "C1478"])

# 三十四.递延收益
fill_second_table_three(49, 1, 6, 0, 7, 1, 7, 2, ["A687", "A688", "A689", "A690", "A691",
                                                  "B687", "B688", "B689", "B690", "B691", "B692",
                                                  "C687", "C688", "C689", "C690", "C691", "C692"])

fill_second_table_three2(49, 1, 7, 3, 7, 4, 6, 5, ["D687", "D688", "D689", "D690", "D691", "D692",
                                                   "E687", "E688", "E689", "E690", "E691", "E692",
                                                   "F687", "F688", "F689", "F690", "F691"])

# 政府补助项目情况
fill_second_table_three(50, 1, 7, 0, 8, 1, 8, 2, ["A695", "A696", "A697", "A698", "A699", "A700",
                                                  "B695", "B696", "B697", "B698", "B699", "B700", "B701",
                                                  "C695", "C696", "C697", "C698", "C699", "C700", "C701"])
fill_second_table_two(50, 1, 8, 3, 8, 4, ["D695", "D696", "D697", "D698", "D699", "D700", "D701",
                                          "E695", "E696", "E697", "E698", "E699", "E700", "E701", ])
fill_second_table_two(50, 1, 8, 5, 8, 6, ["F695", "F696", "F697", "F698", "F699", "F700", "F701",
                                          "G695", "G696", "G697", "G698", "G699", "G700", "G701"])

# 三十五.非流动资产
# 不需要填写

# 三十六.实收资本
# 拆分成3列和3列
# 取数部分，
fill_second_table_two2(52, 2, 9, 0, 10, 1,  ["A709", "A710", "A711", "A712", "A713", "A714", "A715",
                                             "B709", "B710", "B711", "B712", "B713", "B714", "B715", "B716"])
# 取比例部分
fill_second_table_one_per(52, 2, 9, 2, ["L709", "L710", "L711", "L712", "L713", "L714", "L715"])

# 取数部分
fill_second_table_two(52, 2, 10, 3, 10, 4,  ["G709", "G710", "G711", "G712", "G713", "G714", "G715", "G716",
                                             "H709", "H710", "H711", "H712", "H713", "H714", "H715", "H716"])
# 取比例部分
fill_second_table_one_per(52, 2, 9, 5, ["M709", "M710", "M711", "M712", "M713", "M714", "M715"])

# 三十七.资本公积
fill_second_table_two(53, 1, 6, 1, 6, 2, ["B731", "B732", "B733", "B734", "B735",
                                          "C731", "C732", "C733", "C734", "C735"])
fill_second_table_two(53, 1, 6, 3, 6, 4, ["D731", "D732", "D733", "D734", "D735",
                                          "E731", "E732", "E733", "E734", "E735"])

# 三十八.盈余公积
fill_first_table_two(54, 1, 4, 1, 4, 2, ["C1518", "C1519", "C1524",
                                         "C1526", "C1527", "C1532"])
fill_first_table_two(54, 1, 4, 3, 4, 4, ["C1534", "C1535", "C1540",
                                         "C1542", "C1543", "C1548"])

# 三十九.未分配利润
fill_first_table_one(55, 1, 6, 1, ["C1552", "C1553", "C1554", "C1555", "C1555"])
# file_notes.tables[55].cell(8, 1).text = standard_num(str(ws_mid["AA1556"].value))
# file_notes.tables[55].cell(9, 1).text = standard_num(str(ws_mid["C1558"].value))
# file_notes.tables[55].cell(11, 1).text = standard_num(str(ws_mid["c1560"].value))

modify_font(55, 8, 1, standard_num(str(ws_mid["AA1556"].value)))
modify_font(55, 9, 1, standard_num(str(ws_mid["C1558"].value)))
modify_font(55, 11, 1, standard_num(str(ws_mid["c1560"].value)))


# 四十.营业收入和营业成本
fill_first_table_two(56, 1, 4, 1, 4, 2, ["C1565", "C1574", "C1580",
                                         "C1600", "C1609", "C1615"])

# 四十一.税金及附加
fill_first_table_two2(57, 1, 9, 0, 10, 1, ["A1637", "A1638", "A1639", "A1640", "A1641", "A1642", "A1643", "A1644",
                                           "C1637", "C1638", "C1639", "C1640", "C1641", "C1642", "C1643", "C1644",
                                           "C1648"])

# 四十二.销售费用
fill_second_table_two2(58, 1, 21, 0, 22, 1,
                       ["A762", "A763", "A764", "A765", "A766", "A767", "A768", "A769", "A770", "A771",
                        "A772", "A773", "A774", "A775", "A776", "A777", "A778", "A779", "A780", "A781",
                        "B762", "B763", "B764", "B765", "B766", "B767", "B768",
                        "B769", "B770", "B771", "B772", "B773", "B774", "B775",
                        "B776", "B777", "B778", "B779", "B780", "B781", "B782"])

# 四十三.管理费用
fill_second_table_two2(59, 1, 21, 0, 22, 1,
                      ["A786", "A787", "A788", "A789", "A790", "A791", "A792", "A793", "A794", "A795",
                       "A796", "A797", "A798", "A799", "A800", "A801", "A802", "A803", "A804", "A805",
                       "B786", "B787", "B788", "B789", "B790", "B791", "B792", "B793", "B794", "B795",
                       "B796", "B797", "B798", "B799", "B800", "B801", "B802", "B803", "B804", "B805",
                       "B806"])

# 四十四.财务费用
fill_first_table_one(61, 1, 8, 1, ["C1670", "C1671", "C1672", "C1673", "C1674", "C1675", "C1676"])

# 四十五.资产减值损失
fill_first_table_two(62, 1, 17, 1, 17, 2,
                     ["C1688", "C1689", "C1690", "C1691", "C1692", "C1693", "C1694", "C1695", "C1696",
                      "C1697", "C1698", "C16999", "C1700", "C1701", "C1702", "C1703", "C1704", "C1705",
                      "C1706", "C1707", "C1708", "C1709", "C1710", "C1711", "C1712", "C1713", "C1714",
                      "C1715", "C1716", "C1717", "C1718", "C1719", "C1720"])

# 四十六.公允价值变动收益
fill_first_table_two(63, 1, 7, 1, 7, 2, ["C1730", "C1731", "C1732", "C1733", "C1734", "C1735",
                                         "C1737", "C1738", "C1739", "C1740", "C1741", "C1742"])

# 四十七.投资收益
fill_first_table_two(64, 1, 12, 1, 12, 2, ["C1746", "C1747", "C1748", "C1749", "C1750", "C1751",
                                           "C1752", "C1753", "C1754", "C1755", "C1756",
                                           "C1758", "C1759", "C1760", "C1761", "C1762", "C1763",
                                           "C1764", "C1765", "C1766", "C1767", "C1768"])

# 四十八.营业外收入
fill_first_table_two2(65, 1, 9, 0, 10, 1, ["A1816", "A1817", "A1818", "A1819", "A1820", "A1821", "A1822", "A1823",
                                           "C1816", "C1817", "C1818", "C1819", "C1820", "C1821", "C1822", "C1823",
                                           "C1824"])

# 四十九.营业外支出
fill_first_table_two2(66, 1, 6, 0, 7, 1, ["A1849", "A1850", "A1851", "A1852", "A1855",
                                          "C1849", "C1850", "C1851", "C1852", "C1855", "C1856"])

# 五十.所得税费用
fill_first_table_one(67, 1, 4, 1, ["C1879", "C1880", "C1882"])


# file_notes.save(r"D:\data\test\科目余额表报表\1\附注1.10.docx")
file_notes.save(r"D:\work\4输出附注.xlsx")