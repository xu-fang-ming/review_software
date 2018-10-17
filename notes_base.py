import docx
from openpyxl import load_workbook
import re

# 获取文档对象
file_path_notes = r"D:\data\附注\报告-中瑞诚0910.docx"
file_notes = docx.Document(file_path_notes)

# print("表格数:" + str(len(file.tables)))
#
# num = 0
# for i in file_notes.tables:
#     print(str(num)+i.cell(1,0).text)
#     num += 1

# print(file_notes.tables[7].cell(1,1).text)
# print("---------------")
# print(file_notes.tables[7].cell(1,2).text)
# 第二次找到中间表
file_path_mid = r"D:\data\中间表\输出中间表8.xlsx"
wb_mid = load_workbook(filename=file_path_mid, data_only=True)
sheets_mid = wb_mid.sheetnames
sheet_first_mid = sheets_mid[0]  # 中间表1
ws_mid = wb_mid[sheet_first_mid]  # 中间表工作区1

# 中间附注表二
sheet_second_mid = sheets_mid[1]
ws_mid_second = wb_mid[sheet_second_mid]


# ######################开始往附注中填数###########################

def standard_num(x):
    """
    :param x:需要标准化的数字
    :return: 标准化后的数字
    """
    if x == "None":
        return "0.00"
    else:
        if '.' in x:
            num = len(x) % 3
            if num == 0:
                n = x[0:-3]
            elif num == 1:
                n = "00" + x[0:-3]
            elif num == 2:
                n = "0" + x[0:-3]
            lis = re.findall(r'.{3}', n)
            c = ','.join(lis)

            c = c + x[-3:]
        else:
            num = len(x) % 3
            if num == 0:
                n = x
            elif num == 1:
                n = "00" + x
            elif num == 2:
                n = "0" + x

            lis = re.findall(r'.{3}', n)
            c = ','.join(lis)
            c = c + ".00"
        return c.lstrip("0")


def fill_table_two(table_index, start_line, first_column_num,
                   first_column_index, second_column_num, second_column_index,
                   list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param second_column:表格中需要填的第二列
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = standard_num(
            str(ws_mid[list_cell[i - start_line]].value))
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid[list_cell[first_column_num - 1 + (v - start_line)]].value))


def fill_table_one(table_index, start_line, first_column_num,
                   first_column_index, second_column_num, second_column_index,
                   list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param second_column:表格中需要填的第二列
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = standard_num(
            str(ws_mid_second[list_cell[i - start_line]].value))
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column_num - 1 + (v - start_line)]].value))


def fill_table_three(table_index, start_line, first_column, first_index,
                     second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param second_column:表格中需要填的第二列
    :param three_column:表格中需要填的第三列
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        file_notes.tables[table_index].cell(i, first_index).text = standard_num(
            str(ws_mid_second[list_cell[i - start_line]].value))
    for v in range(start_line, second_column):
        file_notes.tables[table_index].cell(v, second_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + (v - start_line)]].value))
    for n in range(start_line, three_column):
        file_notes.tables[table_index].cell(n, three_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + second_column - 1 + (n - start_line)]].value))


def fill_table_three_one(table_index, start_line, first_column, first_index,
                         second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param second_column:表格中需要填的第二列
    :param three_column:表格中需要填的第三列
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        file_notes.tables[table_index].cell(i, first_index).text = standard_num(
            str(ws_mid[list_cell[i - start_line]].value))
    for v in range(start_line, second_column):
        file_notes.tables[table_index].cell(v, second_index).text = standard_num(
            str(ws_mid[list_cell[first_column - 1 + (v - start_line)]].value))
    for n in range(start_line, three_column):
        file_notes.tables[table_index].cell(n, three_index).text = standard_num(
            str(ws_mid[list_cell[first_column - 1 + second_column - 1 + (n - start_line)]].value))


# 一.货币资金
hb1 = standard_num(str(ws_mid["C8"].value))
hb2 = standard_num(str(ws_mid["C9"].value))
hb3 = standard_num(str(ws_mid["C10"].value))
hb4 = standard_num(str(ws_mid["C11"].value))
hb5 = standard_num(str(ws_mid["C14"].value))
hb6 = standard_num(str(ws_mid["C15"].value))
hb7 = standard_num(str(ws_mid["C16"].value))
hb8 = standard_num(str(ws_mid["C17"].value))

file_notes.tables[2].cell(1, 1).text = hb1
file_notes.tables[2].cell(2, 1).text = hb2
file_notes.tables[2].cell(3, 1).text = hb3
file_notes.tables[2].cell(4, 1).text = hb4
file_notes.tables[2].cell(1, 2).text = hb5
file_notes.tables[2].cell(2, 2).text = hb6
file_notes.tables[2].cell(3, 2).text = hb7
file_notes.tables[2].cell(4, 2).text = hb8

# fill_table_two(2, 1, 5, 1, 5, 2, ["C8", "C9", "C10", "C11", "C14", "C15", "C16", "C17"])

# file_notes.save(r"D:\data\附注\4.docx")

# 二.以公允价值计量且其变动计入当期损益的金融资产
gy1 = standard_num(str(ws_mid["C23"].value))
gy2 = standard_num(str(ws_mid["C24"].value))
gy3 = standard_num(str(ws_mid["C25"].value))
gy4 = standard_num(str(ws_mid["C28"].value))
gy5 = standard_num(str(ws_mid["C32"].value))
gy6 = standard_num(str(ws_mid["C34"].value))
gy7 = standard_num(str(ws_mid["C35"].value))
gy8 = standard_num(str(ws_mid["C36"].value))
gy9 = standard_num(str(ws_mid["C39"].value))
gy10 = standard_num(str(ws_mid["C43"].value))

file_notes.tables[3].cell(1, 1).text = gy1
file_notes.tables[3].cell(2, 1).text = gy2
file_notes.tables[3].cell(3, 1).text = gy3
file_notes.tables[3].cell(4, 1).text = gy4
file_notes.tables[3].cell(6, 1).text = gy5
file_notes.tables[3].cell(1, 2).text = gy6
file_notes.tables[3].cell(2, 2).text = gy7
file_notes.tables[3].cell(3, 2).text = gy8
file_notes.tables[3].cell(4, 2).text = gy9
file_notes.tables[3].cell(6, 2).text = gy10

# fill_table_two(3, 1, 5, 1, 5, 2, ["C23", "C24", "C25", "C28", "C34", "C35", "C36", "C39"])
# file_notes.tables[3].cell(6, 1).text = standard_num(str(ws_mid["C32"].value))
# file_notes.tables[3].cell(6, 2).text = standard_num(str(ws_mid["C43"].value))


# 三.应收票据及应收账款
pj_zk1 = standard_num(str(ws_mid["C51"].value))
pj_zk2 = standard_num(str(ws_mid["C78"].value))
pj_zk3 = str(float(pj_zk1)+float(pj_zk2))
pj_zk4 = standard_num(str(ws_mid["C55"].value))
pj_zk5 = standard_num(str(ws_mid["C106"].value))
pj_zk6 = str(float(pj_zk4)+float(pj_zk5))

file_notes.tables[4].cell(1, 1).text = pj_zk1
file_notes.tables[4].cell(2, 1).text = pj_zk2
file_notes.tables[4].cell(3, 1).text = pj_zk3
file_notes.tables[4].cell(1, 2).text = pj_zk4
file_notes.tables[4].cell(2, 2).text = pj_zk5
file_notes.tables[4].cell(3, 2).text = pj_zk6

# fill_table_two(4, 1, 3, 1, 3, 2, ["C51", "C78", "C55", "C106"])
# file_notes.tables[4].cell(3, 1).text = standard_num(str(ws_mid["C51"].value)) + standard_num(str(ws_mid["C78"].value))
# file_notes.tables[4].cell(3, 2).text = standard_num(str(ws_mid["C55"].value)) + standard_num(str(ws_mid["C106"].value))

# 3.1应收票据
pj1 = standard_num(str(ws_mid["C49"].value))
pj2 = standard_num(str(ws_mid["C50"].value))
pj3 = standard_num(str(ws_mid["C51"].value))
pj4 = standard_num(str(ws_mid["C53"].value))
pj5 = standard_num(str(ws_mid["C54"].value))
pj6 = standard_num(str(ws_mid["C55"].value))

file_notes.tables[5].cell(1, 1).text = pj1
file_notes.tables[5].cell(2, 1).text = pj2
file_notes.tables[5].cell(3, 1).text = pj3
file_notes.tables[5].cell(1, 2).text = pj4
file_notes.tables[5].cell(2, 2).text = pj5
file_notes.tables[5].cell(3, 2).text = pj6

# fill_table_two(5, 1, 4, 1, 2, 2, ["C49", "C50", "C51", "C53", "C54", "C55"])

# 3.2应收账款(期末数)
zk1 = standard_num(str(ws_mid["C72"].value))
zk2 = standard_num(str(ws_mid["C73"].value))
zk3 = standard_num(str(ws_mid["C74"].value))
zk4 = standard_num(str(ws_mid["C75"].value))
zk5 = standard_num(str(ws_mid["C78"].value))
if float(zk5) == 0:
    zk6 = "0.00"
    zk7 = "0.00"
    zk8 = "0.00"
    zk9 = "0.00"
else:
    zk6 = "%.2f%%" % (float(zk1)/float(zk5))
    zk7 = "%.2f%%" % (float(zk2)/float(zk5))
    zk8 = "%.2f%%" % (float(zk3)/float(zk5))
    zk9 = "%.2f%%" % (float(zk4)/float(zk5))

file_notes.tables[6].cell(2, 1).text = zk1
file_notes.tables[6].cell(3, 1).text = zk2
file_notes.tables[6].cell(4, 1).text = zk3
file_notes.tables[6].cell(5, 1).text = zk4
file_notes.tables[6].cell(6, 1).text = zk5
file_notes.tables[6].cell(2, 2).text = zk6
file_notes.tables[6].cell(3, 2).text = zk7
file_notes.tables[6].cell(4, 2).text = zk8
file_notes.tables[6].cell(5, 2).text = zk9


# 3.2应收账款(期初数)
zk10 = standard_num(str(ws_mid["C100"].value))
zk11 = standard_num(str(ws_mid["C101"].value))
zk12 = standard_num(str(ws_mid["C102"].value))
zk13 = standard_num(str(ws_mid["C103"].value))
zk14 = standard_num(str(ws_mid["C106"].value))

if float(zk14) == 0:
    zk15 = "0.00"
    zk16 = "0.00"
    zk17 = "0.00"
    zk18 = "0.00"
else:
    zk15 = "%.2f%%" % (float(zk10)/float(zk14))
    zk16 = "%.2f%%" % (float(zk11)/float(zk14))
    zk17 = "%.2f%%" % (float(zk12)/float(zk14))
    zk18 = "%.2f%%" % (float(zk13)/float(zk14))

file_notes.tables[6].cell(2, 3).text = zk10
file_notes.tables[6].cell(3, 3).text = zk11
file_notes.tables[6].cell(4, 3).text = zk12
file_notes.tables[6].cell(5, 3).text = zk13
file_notes.tables[6].cell(6, 3).text = zk14
file_notes.tables[6].cell(2, 4).text = zk15
file_notes.tables[6].cell(3, 4).text = zk16
file_notes.tables[6].cell(4, 4).text = zk17
file_notes.tables[6].cell(5, 4).text = zk18

# fill_table_two(6, 2, 6, 1, 6, 3, ["C72", "C73", "C74", "C75", "C78",
#                                   "C100", "C101", "C102", "C103", "C106"])

# # 3.3 期末余额前五名
zk_qm1 = standard_num(str(ws_mid_second["A60"].value))
zk_qm2 = standard_num(str(ws_mid_second["A61"].value))
zk_qm3 = standard_num(str(ws_mid_second["A62"].value))
zk_qm4 = standard_num(str(ws_mid_second["A63"].value))
zk_qm5 = standard_num(str(ws_mid_second["A64"].value))
zk_qm6 = standard_num(str(ws_mid_second["B60"].value))
zk_qm7 = standard_num(str(ws_mid_second["B61"].value))
zk_qm8 = standard_num(str(ws_mid_second["B62"].value))
zk_qm9 = standard_num(str(ws_mid_second["B63"].value))
zk_qm10 = standard_num(str(ws_mid_second["B64"].value))
zk_qm11 = standard_num(str(ws_mid_second["B65"].value))
zk_qm12 = standard_num(str(ws_mid_second["C60"].value))
zk_qm13 = standard_num(str(ws_mid_second["C61"].value))
zk_qm14 = standard_num(str(ws_mid_second["C62"].value))
zk_qm15 = standard_num(str(ws_mid_second["C63"].value))
zk_qm16 = standard_num(str(ws_mid_second["C64"].value))
zk_qm17 = standard_num(str(ws_mid_second["C65"].value))

file_notes.tables[7].cell(1, 0).text = zk_qm1
file_notes.tables[7].cell(2, 0).text = zk_qm2
file_notes.tables[7].cell(3, 0).text = zk_qm3
file_notes.tables[7].cell(4, 0).text = zk_qm4
file_notes.tables[7].cell(5, 0).text = zk_qm5
file_notes.tables[7].cell(1, 1).text = zk_qm6
file_notes.tables[7].cell(2, 1).text = zk_qm7
file_notes.tables[7].cell(3, 1).text = zk_qm8
file_notes.tables[7].cell(4, 1).text = zk_qm9
file_notes.tables[7].cell(5, 1).text = zk_qm10
file_notes.tables[7].cell(6, 1).text = zk_qm11
file_notes.tables[7].cell(1, 2).text = zk_qm12
file_notes.tables[7].cell(2, 2).text = zk_qm13
file_notes.tables[7].cell(3, 2).text = zk_qm14
file_notes.tables[7].cell(4, 2).text = zk_qm15
file_notes.tables[7].cell(5, 2).text = zk_qm16
file_notes.tables[7].cell(6, 2).text = zk_qm17

# fill_table_three(7, 1, 6, 0, 7, 1, 7, 2, ["A60", "A61", "A62", "A63", "A64",
#                                           "B60", "B61", "B62", "B63", "B64", "B65",
#                                           "C60", "C61", "C62", "C63", "C64", "C65"])

# file_notes.save(r"D:\data\附注\2.docx")
#
# 四.预付款项
# 4.1预付款项(期末数)
yf1 = standard_num(str(ws_mid["C126"].value))
yf2 = standard_num(str(ws_mid["C127"].value))
yf3 = standard_num(str(ws_mid["C128"].value))
yf4 = standard_num(str(ws_mid["C129"].value))
yf5 = standard_num(str(ws_mid["C130"].value))
if float(yf5) == 0:
    yf6 = "0.00"
    yf7 = "0.00"
    yf8 = "0.00"
    yf9 = "0.00"
else:
    yf6 = "%.2f%%" % (float(yf1)/float(yf5))
    yf7 = "%.2f%%" % (float(yf2)/float(yf5))
    yf8 = "%.2f%%" % (float(yf3)/float(yf5))
    yf9 = "%.2f%%" % (float(yf4)/float(yf5))

file_notes.tables[8].cell(2, 1).text = yf1
file_notes.tables[8].cell(3, 1).text = yf2
file_notes.tables[8].cell(4, 1).text = yf3
file_notes.tables[8].cell(5, 1).text = yf4
file_notes.tables[8].cell(6, 1).text = yf5
file_notes.tables[8].cell(2, 2).text = yf6
file_notes.tables[8].cell(3, 2).text = yf7
file_notes.tables[8].cell(4, 2).text = yf8
file_notes.tables[8].cell(5, 2).text = yf9

# 4.2预付款项(期初数)
yf10 = standard_num(str(ws_mid["C132"].value))
yf11 = standard_num(str(ws_mid["C133"].value))
yf12 = standard_num(str(ws_mid["C134"].value))
yf13 = standard_num(str(ws_mid["C135"].value))
yf14 = standard_num(str(ws_mid["C136"].value))

if float(yf14) == 0:
    yf15 = "0.00"
    yf16 = "0.00"
    yf17 = "0.00"
    yf18 = "0.00"
else:
    yf15 = "%.2f%%" % (float(yf10)/float(yf14))
    yf16 = "%.2f%%" % (float(yf11)/float(yf14))
    yf17 = "%.2f%%" % (float(yf12)/float(yf14))
    yf18 = "%.2f%%" % (float(yf13)/float(yf14))

file_notes.tables[8].cell(2, 3).text = yf10
file_notes.tables[8].cell(3, 3).text = yf11
file_notes.tables[8].cell(4, 3).text = yf12
file_notes.tables[8].cell(5, 3).text = yf13
file_notes.tables[8].cell(6, 3).text = yf14
file_notes.tables[8].cell(2, 4).text = yf15
file_notes.tables[8].cell(3, 4).text = yf16
file_notes.tables[8].cell(4, 4).text = yf17
file_notes.tables[8].cell(5, 4).text = yf18

# fill_table_two(8, 2, 6, 1, 6, 3, ["C126", "C127", "C128", "C129", "C130",
#                                   "C132", "C133", "C134", "C135", "C136"])

# 4.3 期末余额前五大明细
yf_qm1 = standard_num(str(ws_mid_second["A88"].value))
yf_qm2 = standard_num(str(ws_mid_second["A89"].value))
yf_qm3 = standard_num(str(ws_mid_second["A90"].value))
yf_qm4 = standard_num(str(ws_mid_second["A91"].value))
yf_qm5 = standard_num(str(ws_mid_second["A92"].value))
yf_qm6 = standard_num(str(ws_mid_second["B88"].value))
yf_qm7 = standard_num(str(ws_mid_second["B89"].value))
yf_qm8 = standard_num(str(ws_mid_second["B90"].value))
yf_qm9 = standard_num(str(ws_mid_second["B91"].value))
yf_qm10 = standard_num(str(ws_mid_second["B92"].value))
yf_qm11 = standard_num(str(ws_mid_second["B93"].value))
yf_qm12 = standard_num(str(ws_mid_second["C88"].value))
yf_qm13 = standard_num(str(ws_mid_second["C89"].value))
yf_qm14 = standard_num(str(ws_mid_second["C90"].value))
yf_qm15 = standard_num(str(ws_mid_second["C91"].value))
yf_qm16 = standard_num(str(ws_mid_second["C92"].value))
yf_qm17 = standard_num(str(ws_mid_second["C93"].value))

file_notes.tables[9].cell(1, 0).text = yf_qm1
file_notes.tables[9].cell(2, 0).text = yf_qm2
file_notes.tables[9].cell(3, 0).text = yf_qm3
file_notes.tables[9].cell(4, 0).text = yf_qm4
file_notes.tables[9].cell(5, 0).text = yf_qm5
file_notes.tables[9].cell(1, 1).text = yf_qm6
file_notes.tables[9].cell(2, 1).text = yf_qm7
file_notes.tables[9].cell(3, 1).text = yf_qm8
file_notes.tables[9].cell(4, 1).text = yf_qm9
file_notes.tables[9].cell(5, 1).text = yf_qm10
file_notes.tables[9].cell(6, 1).text = yf_qm11
file_notes.tables[9].cell(1, 2).text = yf_qm12
file_notes.tables[9].cell(2, 2).text = yf_qm13
file_notes.tables[9].cell(3, 2).text = yf_qm14
file_notes.tables[9].cell(4, 2).text = yf_qm15
file_notes.tables[9].cell(5, 2).text = yf_qm16
file_notes.tables[9].cell(6, 2).text = yf_qm17

# fill_table_three(9, 1, 6, 0, 7, 1, 7, 2, ["A88", "A89", "A90", "A91", "A92",
#                                           "B88", "B89", "B90", "B91", "B92", "B93",
#                                           "C88", "C89", "C90", "C91", "C92", "C93"])

# 五.其他应收款
# 5.1
qt_ys1 = standard_num(str(ws_mid["C147"].value))
qt_ys2 = standard_num(str(ws_mid_second["B112"].value))
qt_ys3 = standard_num(str(ws_mid["C179"].value))
qt_ys4 = str(float(qt_ys1)+float(qt_ys2)+float(qt_ys3))
qt_ys5 = standard_num(str(ws_mid["C154"].value))
qt_ys6 = standard_num(str(ws_mid_second["C112"].value))
qt_ys7 = standard_num(str(ws_mid["C207"].value))
qt_ys8 = str(float(qt_ys5)+float(qt_ys6)+float(qt_ys7))

file_notes.tables[10].cell(1, 1).text = qt_ys1
file_notes.tables[10].cell(2, 1).text = qt_ys2
file_notes.tables[10].cell(3, 1).text = qt_ys3
file_notes.tables[10].cell(4, 1).text = qt_ys4
file_notes.tables[10].cell(1, 2).text = qt_ys5
file_notes.tables[10].cell(2, 2).text = qt_ys6
file_notes.tables[10].cell(3, 2).text = qt_ys7
file_notes.tables[10].cell(4, 2).text = qt_ys8

# 5.2应收利息
ys_lx1 = standard_num(str(ws_mid["A142"].value))
ys_lx2 = standard_num(str(ws_mid["A143"].value))
ys_lx3 = standard_num(str(ws_mid["A144"].value))
ys_lx4 = standard_num(str(ws_mid["A145"].value))
ys_lx5 = standard_num(str(ws_mid["A146"].value))
ys_lx6 = standard_num(str(ws_mid["C142"].value))
ys_lx7 = standard_num(str(ws_mid["C143"].value))
ys_lx8 = standard_num(str(ws_mid["C144"].value))
ys_lx9 = standard_num(str(ws_mid["C145"].value))
ys_lx10 = standard_num(str(ws_mid["C146"].value))
ys_lx11 = standard_num(str(ws_mid["C147"].value))
ys_lx12 = standard_num(str(ws_mid["C149"].value))
ys_lx13 = standard_num(str(ws_mid["C150"].value))
ys_lx14 = standard_num(str(ws_mid["C151"].value))
ys_lx15 = standard_num(str(ws_mid["C152"].value))
ys_lx16 = standard_num(str(ws_mid["C153"].value))
ys_lx17 = standard_num(str(ws_mid["C154"].value))


file_notes.tables[11].cell(1, 0).text = ys_lx1
file_notes.tables[11].cell(2, 0).text = ys_lx2
file_notes.tables[11].cell(3, 0).text = ys_lx3
file_notes.tables[11].cell(4, 0).text = ys_lx4
file_notes.tables[11].cell(5, 0).text = ys_lx5
file_notes.tables[11].cell(1, 1).text = ys_lx6
file_notes.tables[11].cell(2, 1).text = ys_lx7
file_notes.tables[11].cell(3, 1).text = ys_lx8
file_notes.tables[11].cell(4, 1).text = ys_lx9
file_notes.tables[11].cell(5, 1).text = ys_lx10
file_notes.tables[11].cell(6, 1).text = ys_lx11
file_notes.tables[11].cell(1, 2).text = ys_lx12
file_notes.tables[11].cell(2, 2).text = ys_lx13
file_notes.tables[11].cell(3, 2).text = ys_lx14
file_notes.tables[11].cell(4, 2).text = ys_lx15
file_notes.tables[11].cell(5, 2).text = ys_lx16
file_notes.tables[11].cell(6, 2).text = ys_lx17



fill_table_three_one(11, 1, 6, 0, 7, 1, 7, 2, ["A142", "A143", "A144", "A145", "A146",
                                              "C142", "C143", "C144", "C145", "C146", "C147",
                                              "C149", "C150", "C151", "C152", "C153", "C154"])

# # 5.3 应收股利
ys_gl1 = standard_num(str(ws_mid_second["A108"].value))
ys_gl2 = standard_num(str(ws_mid_second["A109"].value))
ys_gl3 = standard_num(str(ws_mid_second["A110"].value))
ys_gl4 = standard_num(str(ws_mid_second["A111"].value))
ys_gl5 = standard_num(str(ws_mid_second["B108"].value))
ys_gl6 = standard_num(str(ws_mid_second["B109"].value))
ys_gl7 = standard_num(str(ws_mid_second["B110"].value))
ys_gl8 = standard_num(str(ws_mid_second["B111"].value))
ys_gl9 = standard_num(str(ws_mid_second["B112"].value))
ys_gl10 = standard_num(str(ws_mid_second["C108"].value))
ys_gl11 = standard_num(str(ws_mid_second["C109"].value))
ys_gl12 = standard_num(str(ws_mid_second["C110"].value))
ys_gl13 = standard_num(str(ws_mid_second["C111"].value))
ys_gl14 = standard_num(str(ws_mid_second["C112"].value))


file_notes.tables[12].cell(1, 0).text = ys_gl1
file_notes.tables[12].cell(2, 0).text = ys_gl2
file_notes.tables[12].cell(3, 0).text = ys_gl3
file_notes.tables[12].cell(4, 0).text = ys_gl4
file_notes.tables[12].cell(1, 1).text = ys_gl5
file_notes.tables[12].cell(2, 1).text = ys_gl6
file_notes.tables[12].cell(3, 1).text = ys_gl7
file_notes.tables[12].cell(4, 1).text = ys_gl8
file_notes.tables[12].cell(5, 1).text = ys_gl9
file_notes.tables[12].cell(1, 2).text = ys_gl10
file_notes.tables[12].cell(2, 2).text = ys_gl11
file_notes.tables[12].cell(3, 2).text = ys_gl12
file_notes.tables[12].cell(4, 2).text = ys_gl13
file_notes.tables[12].cell(5, 2).text = ys_gl14
#
# fill_table_three(12, 1, 5, 0, 6, 1, 6, 2, ["A108", "A109", "A110", "A111",
#                                           "B108", "B109", "B110", "B111", "B112",
#                                           "C108", "C109", "C110", "C111", "C112"])

# 5.4 其他应收款
# 5.4.1其他应收款(期末数)
qt_ysk1 = standard_num(str(ws_mid["C173"].value))
qt_ysk2 = standard_num(str(ws_mid["C174"].value))
qt_ysk3 = standard_num(str(ws_mid["C175"].value))
qt_ysk4 = standard_num(str(ws_mid["C176"].value))
qt_ysk5 = standard_num(str(ws_mid["C179"].value))
if float(qt_ysk5) == 0:
    qt_ysk6 = "0.00"
    qt_ysk7 = "0.00"
    qt_ysk8 = "0.00"
    qt_ysk9 = "0.00"
else:
    qt_ysk6 = "%.2f%%" % (float(qt_ysk1)/float(qt_ysk5))
    qt_ysk7 = "%.2f%%" % (float(qt_ysk2)/float(qt_ysk5))
    qt_ysk8 = "%.2f%%" % (float(qt_ysk3)/float(qt_ysk5))
    qt_ysk9 = "%.2f%%" % (float(qt_ysk4)/float(qt_ysk5))

file_notes.tables[13].cell(2, 1).text = qt_ysk1
file_notes.tables[13].cell(3, 1).text = qt_ysk2
file_notes.tables[13].cell(4, 1).text = qt_ysk3
file_notes.tables[13].cell(5, 1).text = qt_ysk4
file_notes.tables[13].cell(6, 1).text = qt_ysk5
file_notes.tables[13].cell(2, 2).text = qt_ysk6
file_notes.tables[13].cell(3, 2).text = qt_ysk7
file_notes.tables[13].cell(4, 2).text = qt_ysk8
file_notes.tables[13].cell(5, 2).text = qt_ysk9

# 5.4.1其他应收款(期初数)
qt_ysk10 = standard_num(str(ws_mid["C201"].value))
qt_ysk11 = standard_num(str(ws_mid["C202"].value))
qt_ysk12 = standard_num(str(ws_mid["C203"].value))
qt_ysk13 = standard_num(str(ws_mid["C204"].value))
qt_ysk14 = standard_num(str(ws_mid["C207"].value))

if float(qt_ysk14) == 0:
    qt_ysk15 = "0.00"
    qt_ysk16 = "0.00"
    qt_ysk17 = "0.00"
    qt_ysk18 = "0.00"
else:
    qt_ysk15 = "%.2f%%" % (float(qt_ysk10)/float(qt_ysk14))
    qt_ysk16 = "%.2f%%" % (float(qt_ysk11)/float(qt_ysk14))
    qt_ysk17 = "%.2f%%" % (float(qt_ysk12)/float(qt_ysk14))
    qt_ysk18 = "%.2f%%" % (float(qt_ysk13)/float(qt_ysk14))

file_notes.tables[13].cell(2, 3).text = qt_ysk10
file_notes.tables[13].cell(3, 3).text = qt_ysk11
file_notes.tables[13].cell(4, 3).text = qt_ysk12
file_notes.tables[13].cell(5, 3).text = qt_ysk13
file_notes.tables[13].cell(6, 3).text = qt_ysk14
file_notes.tables[13].cell(2, 4).text = qt_ysk15
file_notes.tables[13].cell(3, 4).text = qt_ysk16
file_notes.tables[13].cell(4, 4).text = qt_ysk17
file_notes.tables[13].cell(5, 4).text = qt_ysk18


# # 5.5 其他应收款期末余额前五大
qt_qm1 = standard_num(str(ws_mid_second["A173"].value))
qt_qm2 = standard_num(str(ws_mid_second["A174"].value))
qt_qm3 = standard_num(str(ws_mid_second["A175"].value))
qt_qm4 = standard_num(str(ws_mid_second["A176"].value))
qt_qm5 = standard_num(str(ws_mid_second["A177"].value))
qt_qm6 = standard_num(str(ws_mid_second["C173"].value))
qt_qm7 = standard_num(str(ws_mid_second["C174"].value))
qt_qm8 = standard_num(str(ws_mid_second["C175"].value))
qt_qm9 = standard_num(str(ws_mid_second["C176"].value))
qt_qm10 = standard_num(str(ws_mid_second["C177"].value))
qt_qm11 = standard_num(str(ws_mid_second["C178"].value))
qt_qm12 = standard_num(str(ws_mid_second["E173"].value))
qt_qm13 = standard_num(str(ws_mid_second["E174"].value))
qt_qm14 = standard_num(str(ws_mid_second["E175"].value))
qt_qm15 = standard_num(str(ws_mid_second["E176"].value))
qt_qm16 = standard_num(str(ws_mid_second["E177"].value))
qt_qm17 = standard_num(str(ws_mid_second["E178"].value))

file_notes.tables[14].cell(1, 0).text = qt_qm1
file_notes.tables[14].cell(2, 0).text = qt_qm2
file_notes.tables[14].cell(3, 0).text = qt_qm3
file_notes.tables[14].cell(4, 0).text = qt_qm4
file_notes.tables[14].cell(5, 0).text = qt_qm5
file_notes.tables[14].cell(1, 1).text = qt_qm6
file_notes.tables[14].cell(2, 1).text = qt_qm7
file_notes.tables[14].cell(3, 1).text = qt_qm8
file_notes.tables[14].cell(4, 1).text = qt_qm9
file_notes.tables[14].cell(5, 1).text = qt_qm10
file_notes.tables[14].cell(6, 1).text = qt_qm11
file_notes.tables[14].cell(1, 2).text = qt_qm12
file_notes.tables[14].cell(2, 2).text = qt_qm13
file_notes.tables[14].cell(3, 2).text = qt_qm14
file_notes.tables[14].cell(4, 2).text = qt_qm15
file_notes.tables[14].cell(5, 2).text = qt_qm16
file_notes.tables[14].cell(6, 2).text = qt_qm17

# fill_table_three(14, 1, 6, 0, 7, 1, 7, 2, ["A173", "A174", "A175", "A176", "A177",
#                                           "C173", "C174", "C175", "C176", "C177", "C178",
#                                           "E173", "E174", "E175", "E176", "E177", "E178"])


# # 六.存货
file_notes.tables[15].cell(1, 0).text = standard_num(str(ws_mid["A248"].value))
file_notes.tables[15].cell(2, 0).text = standard_num(str(ws_mid["A249"].value))
file_notes.tables[15].cell(3, 0).text = standard_num(str(ws_mid["A250"].value))
file_notes.tables[15].cell(4, 0).text = standard_num(str(ws_mid["A251"].value))
file_notes.tables[15].cell(5, 0).text = standard_num(str(ws_mid["A252"].value))
file_notes.tables[15].cell(6, 0).text = standard_num(str(ws_mid["A253"].value))
file_notes.tables[15].cell(7, 0).text = standard_num(str(ws_mid["A254"].value))
file_notes.tables[15].cell(1, 1).text = standard_num(str(ws_mid["C248"].value))
file_notes.tables[15].cell(2, 1).text = standard_num(str(ws_mid["C249"].value))
file_notes.tables[15].cell(3, 1).text = standard_num(str(ws_mid["C250"].value))
file_notes.tables[15].cell(4, 1).text = standard_num(str(ws_mid["C251"].value))
file_notes.tables[15].cell(5, 1).text = standard_num(str(ws_mid["C252"].value))
file_notes.tables[15].cell(6, 1).text = standard_num(str(ws_mid["C253"].value))
file_notes.tables[15].cell(7, 1).text = standard_num(str(ws_mid["C254"].value))
file_notes.tables[15].cell(8, 1).text = standard_num(str(ws_mid["C255"].value))
file_notes.tables[15].cell(1, 2).text = standard_num(str(ws_mid["C276"].value))
file_notes.tables[15].cell(2, 2).text = standard_num(str(ws_mid["C277"].value))
file_notes.tables[15].cell(3, 2).text = standard_num(str(ws_mid["C278"].value))
file_notes.tables[15].cell(4, 2).text = standard_num(str(ws_mid["C279"].value))
file_notes.tables[15].cell(5, 2).text = standard_num(str(ws_mid["C280"].value))
file_notes.tables[15].cell(6, 2).text = standard_num(str(ws_mid["C281"].value))
file_notes.tables[15].cell(7, 2).text = standard_num(str(ws_mid["C282"].value))
file_notes.tables[15].cell(8, 2).text = standard_num(str(ws_mid["C283"].value))


# fill_table_three_one(15,1,8,0,9,1,9,2,["A248","A249","A250","A251","A252","A253","A254",
#                                        "C248", "C249", "C250", "C251", "C252", "C253", "C254","C255",
#                                        "C276", "C277", "C278", "C279", "C280", "C281", "C282","C283"])


# def fill_table(table_index,start_line,first_column,second_column,three_column,list_cell):
#     """
#     :param table_index: 表格的索引
#     :param start_line: 表格中要填数据的起始行数
#     :param first_column:表格中需要填的第一列
#     :param second_column:表格中需要填的第二列
#     :param three_column:表格中需要填的第三列
#     :param list_cell:中间表的的单元格编号
#     :return:None
#     """
#     for i in range(start_line, first_column):
#         file_notes.tables[table_index].cell(i, 0).text = standard_num(str(ws_mid[list_cell[i-start_line]].value))
#     for v in range(start_line, second_column):
#         file_notes.tables[table_index].cell(v, 0).text = standard_num(str(ws_mid[list_cell[first_column+(v-start_line)]].value))
#     for n in range(start_line, three_column):
#         file_notes.tables[table_index].cell(n, 0).text = standard_num(str(ws_mid[list_cell[first_column+second_column+(n-start_line)]].value))


file_notes.save(r"D:\data\附注\1.docx")