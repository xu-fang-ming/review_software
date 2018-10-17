import docx
from openpyxl import load_workbook
import re

# 获取文档对象
file_path_notes = r"D:\data\附注\报告-中瑞诚0925.docx"
file_notes = docx.Document(file_path_notes)

# 第二次找到中间表
file_path_mid = r"D:\data\test\科目余额表报表\1\输出中间表1.xlsx"
wb_mid = load_workbook(filename=file_path_mid,data_only=True)
sheets_mid = wb_mid.sheetnames
sheet_first_mid = sheets_mid[0]  # 中间表1
ws_mid = wb_mid[sheet_first_mid]  # 中间表工作区1

# 中间附注表二
sheet_second_mid = sheets_mid[1]
ws_mid_second = wb_mid[sheet_second_mid]


# ######################开始往附注中填数###########################

def standard_num(x):
    """
    :param x:需要标准化的数字
    :return: 标准化后的数字
    """
    if x == "None":
        return "0.00"
    elif x == "0" or x == "0.00":
        return "0.00"
    else:
        if '.' in x:
            num = len(x) % 3
            if num == 0:
                n = x[0:-3]
            elif num == 1:
                n = "00" + x[0:-3]
            elif num == 2:
                n = "0" + x[0:-3]
            lis = re.findall(r'.{3}', n)
            c = ','.join(lis)

            c = c + x[-3:]
        else:
            num = len(x) % 3
            if num == 0:
                n = x
            elif num == 1:
                n = "00" + x
            elif num == 2:
                n = "0" + x

            lis = re.findall(r'.{3}', n)
            c = ','.join(lis)
            c = c + ".00"
        return c.lstrip("0")


def fill_first_table_two(table_index, start_line, first_column_num,
                         first_column_index, second_column_num, second_column_index,
                         list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = standard_num(
            str(ws_mid[list_cell[i - start_line]].value))
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid[list_cell[first_column_num - start_line + (v - start_line)]].value))


def fill_first_table_two2(table_index, start_line, first_column_num,
                          first_column_index, second_column_num, second_column_index,
                          list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号的集合
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = \
            str(ws_mid[list_cell[i - start_line]].value)
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid[list_cell[first_column_num - 1 + (v - start_line)]].value))


def fill_second_table_two(table_index, start_line, first_column_num,
                          first_column_index, second_column_num, second_column_index,
                          list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = standard_num(
            str(ws_mid_second[list_cell[i - start_line]].value))
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column_num - 1 + (v - start_line)]].value))


def fill_second_table_two2(table_index, start_line, first_column_num,
                           first_column_index, second_column_num, second_column_index,
                           list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格中需要填的第一列的索引
    :param second_column_num:表格中需要填的第二列
    :param second_column_index:表格中需要填的第二列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = \
            str(ws_mid_second[list_cell[i - start_line]].value)
    for v in range(start_line, second_column_num):
        file_notes.tables[table_index].cell(v, second_column_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column_num - 1 + (v - start_line)]].value))


def fill_second_table_one(table_index, start_line, first_column_num,
                          first_column_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格的列索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = \
            str(ws_mid_second[list_cell[i - start_line]].value)


def fill_first_table_one(table_index, start_line, first_column_num,
                         first_column_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column_num:表格中需要填的第一列
    :param first_column_index:表格的列索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column_num):
        file_notes.tables[table_index].cell(i, first_column_index).text = \
            standard_num(str(ws_mid[list_cell[i - start_line]].value))


def fill_second_table_three(table_index, start_line, first_column, first_index,
                            second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        file_notes.tables[table_index].cell(i, first_index).text = str(ws_mid_second[list_cell[i - start_line]].value)
    for v in range(start_line, second_column):
        file_notes.tables[table_index].cell(v, second_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + (v - start_line)]].value))
    for n in range(start_line, three_column):
        file_notes.tables[table_index].cell(n, three_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + second_column - 1 + (n - start_line)]].value))


def fill_second_table_three2(table_index, start_line, first_column, first_index,
                             second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        file_notes.tables[table_index].cell(i, first_index).text = standard_num(
            str(ws_mid_second[list_cell[i - start_line]].value))
    for v in range(start_line, second_column):
        file_notes.tables[table_index].cell(v, second_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + (v - start_line)]].value))
    for n in range(start_line, three_column):
        file_notes.tables[table_index].cell(n, three_index).text = standard_num(
            str(ws_mid_second[list_cell[first_column - 1 + second_column - 1 + (n - start_line)]].value))


def fill_first_table_three(table_index, start_line, first_column, first_index,
                           second_column, second_index, three_column, three_index, list_cell):
    """
    :param table_index: 表格的索引
    :param start_line: 表格中要填数据的起始行数
    :param first_column:表格中需要填的第一列
    :param first_index:表格中需要填的第一列的索引
    :param second_column:表格中需要填的第二列
    :param second_index:表格中需要填的第二列的索引
    :param three_column:表格中需要填的第三列
    :param three_index:表格中需要填的第三列的索引
    :param list_cell:中间表的的单元格编号
    :return:None
    """
    for i in range(start_line, first_column):
        file_notes.tables[table_index].cell(i, first_index).text = str(ws_mid[list_cell[i - start_line]].value)
    for v in range(start_line, second_column):
        file_notes.tables[table_index].cell(v, second_index).text = standard_num(
            str(ws_mid[list_cell[first_column - 1 + (v - start_line)]].value))
    for n in range(start_line, three_column):
        file_notes.tables[table_index].cell(n, three_index).text = standard_num(
            str(ws_mid[list_cell[first_column - 1 + second_column - 1 + (n - start_line)]].value))


# 一.货币资金
fill_first_table_two(2, 1, 5, 1, 5, 2, ["C8", "C9", "C10", "C11", "C14", "C15", "C16", "C17"])

# 3.2应收账款
# 取数部分
fill_first_table_two(6, 2, 7, 1, 7, 3, ["AA80", "AA81", "AA82", "AA83", "AA86",
                                        "AA108", "AA109", "AA110", "AA111", "AA114"])

# 取比例部分
fill_first_table_two(6, 2, 6, 2, 6, 4, ["AB80", "AB81", "AB82", "AB83",
                                        "AB108", "AB109", "AB110", "AB111"])


file_notes.save(r"D:\data\test\科目余额表报表\1\附注2.docx")